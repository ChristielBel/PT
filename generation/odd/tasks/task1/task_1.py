import os.path
import random
from math import factorial
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as align

import generation.writer as writer

first_val = 20
second_val = 5


def generate_task_1(target_doc_path):
    global first_val
    global second_val

    first_val = random.randint(20, 40)
    second_val = random.randint(5, int(first_val / 2))
    replacement_values = [str(first_val), str(second_val)]
    # Open the source document
    source_doc = Document(os.path.abspath("generation/odd/tasks/task1/Task1.docx"))

    # Initialize index to track replacement value
    replacement_index = 0

    # Clone the source document
    target_doc = Document(target_doc_path)

    # Iterate over paragraphs
    for paragraph in source_doc.paragraphs:
        # Place all values to all paragraphs
        while paragraph.text.find("+") != -1:
            paragraph.text = paragraph.text.replace("+", replacement_values[replacement_index], 1)
            replacement_index += 1

        target_doc.add_paragraph(paragraph.text)

    # Save the modified document
    target_doc.save(target_doc_path)


def calculate_task(target_doc_path):
    global first_val
    global second_val

    ans = factorial(first_val) / factorial(first_val - second_val) * factorial(second_val)

    writer.write_text(target_doc_path, "1. " + str(ans), "Arial", 14, align.RIGHT, False)



