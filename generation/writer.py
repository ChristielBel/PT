import os

import docx
from docx import Document
from docx.shared import Pt


def write_text(doc_path, text, font_name, font_size, alignment, is_bold):
    doc = docx.Document(doc_path)

    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment

    run = paragraph.add_run(text)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = is_bold

    doc.save(doc_path)


def replace_placeholders_and_write_to_target(source_doc_path, target_doc_path, replacement_values, symbol):
    # Open the source document
    source_doc = Document(os.path.abspath(source_doc_path))

    # Initialize index to track replacement value
    replacement_index = 0

    # Clone the source document
    target_doc = Document(target_doc_path)

    # Iterate over paragraphs
    for paragraph in source_doc.paragraphs:
        # Place all values to all paragraphs
        while paragraph.text.find(symbol) != -1:
            paragraph.text = paragraph.text.replace(symbol, str(replacement_values[replacement_index]), 1)

            replacement_index += 1

        new_paragraph = target_doc.add_paragraph()
        run = new_paragraph.add_run(paragraph.text)
        font = run.font
        font.name = "Arial"
        font.size = Pt(14)

    # Save the modified document
    target_doc.save(target_doc_path)

